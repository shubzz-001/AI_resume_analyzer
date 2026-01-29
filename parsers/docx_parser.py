from docx import Document
from pathlib import Path
from typing import Optional, List
from io import BytesIO
import logging

from parsers.base_parser import BaseParser

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):
    """Parser for DOCX documents."""

    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.docx']

    def parse(self, file_path: Path) -> str:

        if not self.validate_file(file_path):
            raise FileNotFoundError(f"Invalid DOCX file: {file_path}")

        try:
            text = self._extract_text_from_docx(file_path)

            if not text or len(text.strip()) < 10:
                raise ValueError("DOCX appears to be empty or contains no extractable text")

            return self.clean_text(text)

        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
            raise ValueError(f"Failed to parse DOCX: {str(e)}")

    def _extract_text_from_docx(self, file_path: Path) -> str:

        try:
            doc = Document(file_path)
            text_parts = []

            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            logger.debug(f"Extracted {len(text_parts)} text blocks from DOCX")
            return "\n".join(text_parts)

        except Exception as e:
            logger.error(f"Error in _extract_text_from_docx: {str(e)}")
            raise

    def parse_from_bytes(self, file_bytes: bytes, filename: str = "resume.docx") -> str:

        try:
            file_stream = BytesIO(file_bytes)
            doc = Document(file_stream)
            text_parts = []

            # Extract from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            text = "\n".join(text_parts)

            if not text or len(text.strip()) < 10:
                raise ValueError("DOCX appears to be empty or contains no extractable text")

            return self.clean_text(text)

        except Exception as e:
            logger.error(f"Error parsing DOCX bytes from {filename}: {str(e)}")
            raise ValueError(f"Failed to parse DOCX: {str(e)}")

    def get_paragraph_count(self, file_path: Path) -> int:

        try:
            doc = Document(file_path)
            return len([p for p in doc.paragraphs if p.text.strip()])
        except Exception as e:
            logger.error(f"Error getting paragraph count: {str(e)}")
            return 0

    def extract_headings(self, file_path: Path) -> List[str]:

        headings = []

        try:
            doc = Document(file_path)

            for para in doc.paragraphs:
                # Check if paragraph is a heading
                if para.style.name.startswith('Heading'):
                    headings.append(para.text.strip())

            logger.debug(f"Found {len(headings)} headings in DOCX")

        except Exception as e:
            logger.warning(f"Could not extract headings: {str(e)}")

        return headings

    def extract_metadata(self, file_path: Path) -> dict:

        metadata = self.get_file_info(file_path)

        try:
            doc = Document(file_path)
            core_props = doc.core_properties

            metadata.update({
                "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),
                "table_count": len(doc.tables),
                "author": core_props.author or "Unknown",
                "title": core_props.title or "Untitled",
                "created": str(core_props.created) if core_props.created else None,
                "modified": str(core_props.modified) if core_props.modified else None
            })

        except Exception as e:
            logger.warning(f"Could not extract DOCX metadata: {str(e)}")

        return metadata


# Convenience function for direct use
def extract_text_from_docx(file_path: Path) -> str:

    parser = DOCXParser()
    return parser.parse(file_path)